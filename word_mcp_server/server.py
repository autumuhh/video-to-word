import asyncio
import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from mcp.server.models import InitializationOptions
import mcp.types as types
from mcp.server import Server
from mcp.server.stdio import stdio_server

# Initialize the server
app = Server("word-mcp-server")

def set_run_font(run, font_name='SimSun', size=12):
    """Sets the font for a run, handling complex scripts (Chinese)."""
    run.font.name = 'Times New Roman'  # For ASCII
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    # Set East Asia font (Chinese)
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def add_formatted_paragraph(doc, text, bold_title=False):
    """
    Parses simple markdown bold (**text**) and applies strict formatting.
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if not part: continue
        
        is_bold = part.startswith('**') and part.endswith('**')
        clean_text = part.strip('*') if is_bold else part
        
        run = p.add_run(clean_text)
        run.bold = is_bold
        
        # Apply Strict Font Settings to EVERY run
        set_run_font(run, font_name='SimSun', size=12)

def generate_docx(content: str, output_path: str, image_map: dict = None):
    """
    Generates a Word document with strict academic formatting.
    """
    doc = Document()
    
    # Configure Normal Style defaults
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Title
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, font_name='SimHei', size=16) # Title often SimHei
                
        elif line.startswith('## '):
            # Heading 1
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                set_run_font(run, font_name='SimHei', size=14)
                
        elif line.startswith('### '):
            # Heading 2
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                set_run_font(run, font_name='SimHei', size=13)
                
        elif '[INSERT_IMAGE:' in line:
            # Image handling remains the same...
            match = re.search(r'\[INSERT_IMAGE:\s*(.*?)\]', line)
            if match:
                img_key = match.group(1)
                if image_map and img_key in image_map:
                    img_path = image_map[img_key]
                    if os.path.exists(img_path):
                        try:
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(6.0))
                            
                            caption = doc.add_paragraph(f"Figure: {img_key}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in caption.runs:
                                set_run_font(run, font_name='Times New Roman', size=10)
                        except:
                            pass
        else:
            # Standard paragraph with strict font parsing
            add_formatted_paragraph(doc, line)
            
    # Ensure directory exists
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path

@app.call_tool()
async def call_tool(name: str, arguments: dict) -> list[types.TextContent | types.ImageContent | types.EmbeddedResource]:
    if name == "generate_word_doc":
        content = arguments.get("content")
        output_path = arguments.get("output_path")
        image_map = arguments.get("image_map", {})
        
        if not content or not output_path:
            raise ValueError("Missing required arguments: content, output_path")
            
        try:
            saved_path = generate_docx(content, output_path, image_map)
            return [types.TextContent(type="text", text=f"Successfully generated Word document at: {saved_path}")]
        except Exception as e:
            return [types.TextContent(type="text", text=f"Error generating document: {str(e)}")]
            
    raise ValueError(f"Tool not found: {name}")

@app.list_tools()
async def list_tools() -> list[types.Tool]:
    return [
        types.Tool(
            name="generate_word_doc",
            description="Generate a Word document (.docx) from Markdown text and images",
            inputSchema={
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string", 
                        "description": "Markdown content. Supports headers (#) and [INSERT_IMAGE: key] placeholders."
                    },
                    "output_path": {
                        "type": "string", 
                        "description": "Absolute path to save the generated .docx file"
                    },
                    "image_map": {
                        "type": "object", 
                        "description": "Dictionary mapping image keys (used in content placeholders) to file paths"
                    }
                },
                "required": ["content", "output_path"]
            }
        )
    ]

async def main():
    async with stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="word-mcp-server",
                server_version="0.1.0",
                capabilities=app.get_capabilities(
                    notification_options=types.NotificationOptions(),
                    experimental_capabilities={},
                ),
            ),
        )

if __name__ == "__main__":
    asyncio.run(main())